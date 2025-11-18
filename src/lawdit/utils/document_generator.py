"""
Document Generation Utilities

Utilities for creating Word documents and HTML dashboards.
"""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


class WordDocumentGenerator:
    """Generator for creating professional Word documents."""

    def __init__(self):
        """Initialize the document generator."""
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles

        # Title style
        if "CustomTitle" not in styles:
            title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)

        # Heading styles are built-in, just configure them
        for level in range(1, 4):
            heading_style = styles[f"Heading {level}"]
            heading_style.font.color.rgb = RGBColor(0, 51, 102)

    def add_cover_page(self, title: str, subtitle: str = None):
        """Add a cover page to the document.

        Args:
            title: Main title
            subtitle: Optional subtitle
        """
        # Add title
        title_para = self.doc.add_paragraph(title)
        title_para.style = "CustomTitle"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add space
        self.doc.add_paragraph()

        # Add subtitle if provided
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.runs[0].font.size = Pt(14)

        # Add date
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break
        self.doc.add_page_break()

    def add_executive_summary(self, summary_text: str):
        """Add an executive summary section.

        Args:
            summary_text: The summary content
        """
        self.doc.add_heading("Executive Summary", level=1)
        self.doc.add_paragraph(summary_text)
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add a placeholder for table of contents.

        Note: Word table of contents must be generated within Word itself.
        """
        self.doc.add_heading("Table of Contents", level=1)
        self.doc.add_paragraph(
            "To generate the table of contents:\n"
            "1. Click here\n"
            "2. Go to References > Table of Contents\n"
            "3. Choose a style"
        )
        self.doc.add_page_break()

    def add_risk_section(
        self, category: str, risks: List[Dict[str, Any]], overview: str = None
    ):
        """Add a risk category section.

        Args:
            category: Risk category name
            risks: List of risk dictionaries
            overview: Optional category overview
        """
        self.doc.add_heading(category, level=1)

        if overview:
            self.doc.add_paragraph(overview)
            self.doc.add_paragraph()

        for idx, risk in enumerate(risks, 1):
            # Risk title
            self.doc.add_heading(risk.get("title", f"Risk {idx}"), level=2)

            # Severity
            severity = risk.get("severity", "Unknown")
            severity_para = self.doc.add_paragraph(f"Severity: ")
            severity_run = severity_para.add_run(severity)
            severity_run.bold = True

            # Color code by severity
            if severity.lower() == "critical":
                severity_run.font.color.rgb = RGBColor(192, 0, 0)
            elif severity.lower() == "high":
                severity_run.font.color.rgb = RGBColor(255, 102, 0)
            elif severity.lower() == "medium":
                severity_run.font.color.rgb = RGBColor(255, 192, 0)

            # Description
            if "description" in risk:
                self.doc.add_heading("Description", level=3)
                self.doc.add_paragraph(risk["description"])

            # Supporting Evidence
            if "evidence" in risk:
                self.doc.add_heading("Supporting Evidence", level=3)
                self.doc.add_paragraph(risk["evidence"])

            # Impact
            if "impact" in risk:
                self.doc.add_heading("Potential Impact", level=3)
                self.doc.add_paragraph(risk["impact"])

            # Recommendations
            if "recommendations" in risk:
                self.doc.add_heading("Recommendations", level=3)
                self.doc.add_paragraph(risk["recommendations"])

            self.doc.add_paragraph()  # Spacing

    def add_risk_matrix_table(self, risks: List[Dict[str, Any]]):
        """Add a summary risk matrix table.

        Args:
            risks: List of all risks
        """
        self.doc.add_heading("Risk Matrix Summary", level=2)

        # Create table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Risk"
        header_cells[1].text = "Category"
        header_cells[2].text = "Severity"
        header_cells[3].text = "Documents"

        # Add risks
        for risk in risks:
            row_cells = table.add_row().cells
            row_cells[0].text = risk.get("title", "")
            row_cells[1].text = risk.get("category", "")
            row_cells[2].text = risk.get("severity", "")
            row_cells[3].text = risk.get("documents", "")

    def save(self, output_path: str):
        """Save the document to a file.

        Args:
            output_path: Path to save the document
        """
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)


class DashboardGenerator:
    """Generator for creating interactive HTML dashboards."""

    def __init__(self):
        """Initialize the dashboard generator."""
        self.risks: List[Dict[str, Any]] = []
        self.categories: List[str] = []
        self.severities: List[str] = ["Critical", "High", "Medium", "Low"]

    def add_risk(self, risk: Dict[str, Any]):
        """Add a risk to the dashboard.

        Args:
            risk: Risk dictionary with title, category, severity, etc.
        """
        self.risks.append(risk)
        if risk.get("category") and risk["category"] not in self.categories:
            self.categories.append(risk["category"])

    def generate_html(self) -> str:
        """Generate the complete HTML dashboard.

        Returns:
            HTML string for the dashboard
        """
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Legal Risk Analysis Dashboard</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: #f5f7fa;
            padding: 20px;
        }}

        .container {{
            max-width: 1400px;
            margin: 0 auto;
        }}

        h1 {{
            color: #2c3e50;
            margin-bottom: 30px;
        }}

        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }}

        .stat-card {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}

        .stat-value {{
            font-size: 2em;
            font-weight: bold;
            color: #3498db;
        }}

        .stat-label {{
            color: #7f8c8d;
            margin-top: 5px;
        }}

        .filters {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }}

        .filter-group {{
            margin-bottom: 15px;
        }}

        .filter-group label {{
            display: block;
            margin-bottom: 5px;
            font-weight: 500;
            color: #2c3e50;
        }}

        select, input {{
            width: 100%;
            padding: 8px;
            border: 1px solid #ddd;
            border-radius: 4px;
        }}

        .risks-grid {{
            display: grid;
            gap: 15px;
        }}

        .risk-card {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            border-left: 4px solid #3498db;
        }}

        .risk-card.critical {{
            border-left-color: #c0392b;
        }}

        .risk-card.high {{
            border-left-color: #e67e22;
        }}

        .risk-card.medium {{
            border-left-color: #f39c12;
        }}

        .risk-card.low {{
            border-left-color: #27ae60;
        }}

        .risk-header {{
            display: flex;
            justify-content: space-between;
            align-items: start;
            margin-bottom: 10px;
        }}

        .risk-title {{
            font-size: 1.2em;
            font-weight: 600;
            color: #2c3e50;
        }}

        .risk-severity {{
            padding: 4px 12px;
            border-radius: 4px;
            font-size: 0.85em;
            font-weight: 600;
            color: white;
        }}

        .severity-critical {{
            background: #c0392b;
        }}

        .severity-high {{
            background: #e67e22;
        }}

        .severity-medium {{
            background: #f39c12;
        }}

        .severity-low {{
            background: #27ae60;
        }}

        .risk-category {{
            color: #7f8c8d;
            margin-bottom: 10px;
        }}

        .risk-description {{
            color: #34495e;
            line-height: 1.6;
        }}

        .hidden {{
            display: none;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Legal Risk Analysis Dashboard</h1>

        <div class="stats">
            <div class="stat-card">
                <div class="stat-value" id="total-risks">{len(self.risks)}</div>
                <div class="stat-label">Total Risks Identified</div>
            </div>
            <div class="stat-card">
                <div class="stat-value" id="critical-count">{self._count_by_severity('Critical')}</div>
                <div class="stat-label">Critical Risks</div>
            </div>
            <div class="stat-card">
                <div class="stat-value" id="high-count">{self._count_by_severity('High')}</div>
                <div class="stat-label">High Risks</div>
            </div>
            <div class="stat-card">
                <div class="stat-value" id="category-count">{len(self.categories)}</div>
                <div class="stat-label">Risk Categories</div>
            </div>
        </div>

        <div class="filters">
            <div class="filter-group">
                <label for="category-filter">Filter by Category:</label>
                <select id="category-filter" onchange="filterRisks()">
                    <option value="">All Categories</option>
                    {''.join(f'<option value="{cat}">{cat}</option>' for cat in self.categories)}
                </select>
            </div>
            <div class="filter-group">
                <label for="severity-filter">Filter by Severity:</label>
                <select id="severity-filter" onchange="filterRisks()">
                    <option value="">All Severities</option>
                    {''.join(f'<option value="{sev}">{sev}</option>' for sev in self.severities)}
                </select>
            </div>
            <div class="filter-group">
                <label for="search">Search:</label>
                <input type="text" id="search" placeholder="Search risks..." oninput="filterRisks()">
            </div>
        </div>

        <div class="risks-grid" id="risks-container">
            {self._generate_risk_cards()}
        </div>
    </div>

    <script>
        const risks = {self._risks_to_json()};

        function filterRisks() {{
            const category = document.getElementById('category-filter').value.toLowerCase();
            const severity = document.getElementById('severity-filter').value.toLowerCase();
            const search = document.getElementById('search').value.toLowerCase();

            const cards = document.querySelectorAll('.risk-card');

            cards.forEach(card => {{
                const cardCategory = card.dataset.category.toLowerCase();
                const cardSeverity = card.dataset.severity.toLowerCase();
                const cardText = card.textContent.toLowerCase();

                const categoryMatch = !category || cardCategory === category;
                const severityMatch = !severity || cardSeverity === severity;
                const searchMatch = !search || cardText.includes(search);

                if (categoryMatch && severityMatch && searchMatch) {{
                    card.classList.remove('hidden');
                }} else {{
                    card.classList.add('hidden');
                }}
            }});
        }}
    </script>
</body>
</html>"""
        return html

    def _count_by_severity(self, severity: str) -> int:
        """Count risks by severity level."""
        return sum(1 for r in self.risks if r.get("severity") == severity)

    def _generate_risk_cards(self) -> str:
        """Generate HTML for all risk cards."""
        cards = []
        for risk in self.risks:
            severity = risk.get("severity", "Unknown").lower()
            cards.append(f"""
            <div class="risk-card {severity}" data-category="{risk.get('category', '')}" data-severity="{risk.get('severity', '')}">
                <div class="risk-header">
                    <div class="risk-title">{risk.get('title', 'Untitled Risk')}</div>
                    <span class="risk-severity severity-{severity}">{risk.get('severity', 'Unknown')}</span>
                </div>
                <div class="risk-category">{risk.get('category', 'Uncategorized')}</div>
                <div class="risk-description">{risk.get('description', 'No description available.')}</div>
            </div>
            """)
        return '\n'.join(cards)

    def _risks_to_json(self) -> str:
        """Convert risks to JSON for JavaScript."""
        import json

        return json.dumps(self.risks)

    def save(self, output_path: str):
        """Save the dashboard to an HTML file.

        Args:
            output_path: Path to save the HTML file
        """
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        html = self.generate_html()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
